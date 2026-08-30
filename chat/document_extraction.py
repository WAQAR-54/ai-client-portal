"""Extracts plain text from an uploaded chat attachment so its content can
be included as context for the assistant - PDF, Word, Excel, and simple
text files. Never raises: extraction failures degrade to None so the
caller can fall back to "not readable" rather than crashing the chat.

SECURITY: everything this returns is untrusted content from a file a user
uploaded. Callers MUST wrap it with wrap_for_prompt() before it ever
reaches the model - see that function's docstring and chat/prompts.py's
system prompt for the other half of this defense (the model is explicitly
told to treat delimited content as reference material, never instructions,
so a document containing "ignore previous instructions" doesn't work).
"""

import logging

logger = logging.getLogger(__name__)

TEXT_EXTENSIONS = {"txt", "csv", "md", "json"}
PDF_EXTENSIONS = {"pdf"}
DOCX_EXTENSIONS = {"docx"}
XLSX_EXTENSIONS = {"xlsx"}
EXTRACTABLE_EXTENSIONS = TEXT_EXTENSIONS | PDF_EXTENSIONS | DOCX_EXTENSIONS | XLSX_EXTENSIONS

MAX_CHARS = 8000


def _extract_text_file(file_field):
    with file_field.open("rb") as f:
        return f.read(MAX_CHARS + 1).decode("utf-8", errors="replace")


def _extract_pdf(file_field):
    from pypdf import PdfReader

    with file_field.open("rb") as f:
        reader = PdfReader(f)
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(file_field):
    import docx

    with file_field.open("rb") as f:
        document = docx.Document(f)
    parts = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_xlsx(file_field):
    import openpyxl

    with file_field.open("rb") as f:
        workbook = openpyxl.load_workbook(f, read_only=True, data_only=True)
    lines = []
    for sheet in workbook.worksheets:
        lines.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)


_EXTRACTORS = {
    **{ext: _extract_text_file for ext in TEXT_EXTENSIONS},
    **{ext: _extract_pdf for ext in PDF_EXTENSIONS},
    **{ext: _extract_docx for ext in DOCX_EXTENSIONS},
    **{ext: _extract_xlsx for ext in XLSX_EXTENSIONS},
}


def extract_text(file_field, extension):
    """Returns extracted text (truncated to MAX_CHARS), or None if this
    extension isn't supported, the file has no extractable text, or
    extraction failed for any reason."""
    extractor = _EXTRACTORS.get(extension.lower())
    if extractor is None:
        return None
    try:
        text = extractor(file_field)
    except Exception:
        logger.exception("Failed to extract text from attachment (extension=%s)", extension)
        return None
    text = (text or "").strip()
    if not text:
        return None
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + "\n[...truncated...]"
    return text


def wrap_for_prompt(filename, text):
    """Delimits extracted document text so it reads as clearly-marked
    reference material, not as part of the conversation. The delimiter
    alone isn't the defense - it only works paired with the system
    prompt's explicit instruction (chat/prompts.py) to never treat
    anything inside these markers as instructions."""
    return f"[BEGIN ATTACHED DOCUMENT: {filename}]\n{text}\n[END ATTACHED DOCUMENT: {filename}]"
